"""Parser for DOCX files (TRU Moodle FAQ and similar)."""

import logging
from pathlib import Path

from docx import Document as DocxDocument

from ...models.schemas import Document

logger = logging.getLogger(__name__)


def parse_docx(file_path: Path, source_name: str = "tru_faq") -> list[Document]:
    """Parse a DOCX file into documents, splitting by headings.

    Each heading-delimited section becomes a separate Document.
    If no headings are found, the entire file becomes one Document.
    """
    try:
        doc = DocxDocument(str(file_path))
    except Exception as e:
        logger.error(f"Failed to open DOCX {file_path}: {e}")
        return []

    sections: list[Document] = []
    current_title = file_path.stem  # Default title from filename
    current_text_parts: list[str] = []
    heading_hierarchy: list[str] = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if not text:
            continue

        # Detect headings
        if style_name.startswith("Heading"):
            # Save the previous section if it has content
            if current_text_parts:
                section_text = "\n".join(current_text_parts)
                if len(section_text) >= 30:
                    sections.append(Document(
                        source=source_name,
                        title=current_title,
                        text=section_text,
                        slug=_slugify(current_title),
                        categories=["FAQ"] if source_name == "tru_faq" else [],
                    ))
                current_text_parts = []

            # Update heading hierarchy
            try:
                level = int(style_name.replace("Heading ", "").replace("Heading", "1"))
            except ValueError:
                level = 1

            # Trim hierarchy to current level
            heading_hierarchy = heading_hierarchy[:level - 1]
            heading_hierarchy.append(text)
            current_title = " > ".join(heading_hierarchy) if heading_hierarchy else text

        else:
            current_text_parts.append(text)

    # Don't forget the last section
    if current_text_parts:
        section_text = "\n".join(current_text_parts)
        if len(section_text) >= 30:
            sections.append(Document(
                source=source_name,
                title=current_title,
                text=section_text,
                slug=_slugify(current_title),
                categories=["FAQ"] if source_name == "tru_faq" else [],
            ))

    # If no sections were created (no headings found), create one from the whole doc
    if not sections:
        full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        if full_text:
            sections.append(Document(
                source=source_name,
                title=file_path.stem,
                text=full_text,
                slug=_slugify(file_path.stem),
                categories=["FAQ"] if source_name == "tru_faq" else [],
            ))

    logger.info(f"Parsed {len(sections)} sections from {file_path.name}")
    return sections


def _slugify(text: str) -> str:
    """Create a simple slug from text."""
    import re
    slug = text.lower()
    slug = re.sub(r"[^a-z0-9\s-]", "", slug)
    slug = re.sub(r"[\s]+", "-", slug)
    slug = slug.strip("-")
    return slug[:80]  # Limit length
